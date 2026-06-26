import docx
import re
from typing import Dict, List, Any

def extract_text_from_docx(file_bytes: bytes) -> str:
    from io import BytesIO
    doc = docx.Document(BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text)

def parse_jd_requirements(text: str) -> Dict[str, Any]:
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    
    extracted = {
        "title": "Senior AI Engineer",
        "extracted_skills": [],
        "extracted_experience": "",
        "extracted_education": "",
        "extracted_certifications": [],
        "extracted_responsibilities": []
    }
    
    # 1. Parse job title
    title_match = re.search(r"(?:Job Title|Role|Position|Title):\s*(.*)", text, re.IGNORECASE)
    if title_match:
        extracted["title"] = title_match.group(1).strip()
    else:
        # Fallback to the first line if it looks like a title
        if lines and len(lines[0]) < 60 and not any(w in lines[0].lower() for w in ["hiring", "looking for", "description", "we are"]):
            extracted["title"] = lines[0].strip()
            
    # 2. Broad skills vocabulary dictionary matching (Frontend, Backend, DevOps, ML/AI, QA, Mobile, PM)
    skills_vocabulary = {
        # AI / ML
        "python": "Python", "pytorch": "PyTorch", "tensorflow": "TensorFlow", "embeddings": "Embeddings",
        "vector search": "Vector Search", "semantic search": "Semantic Search", "rag": "RAG", "llm": "LLM",
        "fine-tuning": "Fine-Tuning", "lora": "LoRA", "qlora": "QLoRA", "peft": "PEFT", "xgboost": "XGBoost",
        "scikit-learn": "Scikit-Learn", "ndcg": "NDCG", "mrr": "MRR", "map": "MAP", "a/b testing": "A/B Testing",
        "pinecone": "Pinecone", "weaviate": "Weaviate", "qdrant": "Qdrant", "milvus": "Milvus", "faiss": "FAISS",
        "elasticsearch": "Elasticsearch", "opensearch": "Opensearch", "pandas": "Pandas", "numpy": "NumPy",
        "spacy": "spaCy", "nltk": "NLTK", "keras": "Keras",
        # Java & Backend
        "java": "Java", "spring boot": "Spring Boot", "spring": "Spring", "hibernate": "Hibernate",
        "microservices": "Microservices", "rest api": "REST API", "restful api": "RESTful API", "jpa": "JPA",
        "junit": "JUnit", "maven": "Maven", "gradle": "Gradle",
        # Javascript & Frontend
        "javascript": "JavaScript", "typescript": "TypeScript", "react": "React", "angular": "Angular",
        "vue": "Vue", "next.js": "Next.js", "nextjs": "Next.js", "html": "HTML", "css": "CSS",
        "tailwind": "Tailwind CSS", "bootstrap": "Bootstrap", "sass": "Sass", "webpack": "Webpack",
        "vite": "Vite", "nodejs": "Node.js", "node.js": "Node.js", "express": "Express",
        "nestjs": "NestJS", "redux": "Redux", "svelte": "Svelte",
        # DevOps & Cloud
        "aws": "AWS", "gcp": "GCP", "azure": "Azure", "docker": "Docker", "kubernetes": "Kubernetes",
        "terraform": "Terraform", "ansible": "Ansible", "jenkins": "Jenkins", "ci/cd": "CI/CD",
        "git": "Git", "github": "GitHub", "gitlab": "GitLab", "linux": "Linux", "bash": "Bash",
        # Other Languages
        "golang": "Go", "go": "Go", "rust": "Rust", "c++": "C++", "c#": "C#", ".net": ".NET",
        "ruby": "Ruby", "rails": "Ruby on Rails", "php": "PHP", "kotlin": "Kotlin", "swift": "Swift",
        "scala": "Scala",
        # Databases & Data
        "sql": "SQL", "postgresql": "PostgreSQL", "postgres": "PostgreSQL", "mysql": "MySQL",
        "sqlite": "SQLite", "mongodb": "MongoDB", "redis": "Redis", "cassandra": "Cassandra",
        "dynamodb": "DynamoDB", "snowflake": "Snowflake", "databricks": "Databricks",
        # PM & QA & General
        "scrum": "Scrum", "agile": "Agile", "jira": "Jira", "selenium": "Selenium", "cypress": "Cypress",
        "jest": "Jest", "playwright": "Playwright", "project management": "Project Management",
        "product management": "Product Management"
    }

    found_skills = set()
    for kw, display in skills_vocabulary.items():
        if re.search(r"\b" + re.escape(kw) + r"\b", text, re.IGNORECASE):
            found_skills.add(display)

    # 3. Dynamic Section-based Requirement Extraction
    # Look for requirements/skills sections and extract capitalized words/phrases
    req_section = False
    stop_words_extra = {
        "the", "and", "our", "you", "for", "with", "that", "this", "from", "your", "have", "will", "work",
        "team", "role", "data", "system", "code", "design", "build", "using", "years", "experience", "degree",
        "field", "science", "engineering", "technical", "strong", "knowledge", "ability", "written", "verbal"
    }
    
    for line in lines:
        line_lower = line.lower()
        if any(w in line_lower for w in ["requirement", "qualification", "skills", "experience", "what you need", "what you bring", "must have", "stack", "technologies"]):
            req_section = True
            continue
        if req_section:
            if any(w in line_lower for w in ["responsibility", "responsibilities", "what you will do", "about us", "benefits", "offer"]):
                req_section = False
            elif line.startswith("-") or line.startswith("*") or line.startswith("•") or (len(line) > 5 and line[0].isupper()):
                # Extract capitalized words that could be custom technologies/skills
                # e.g., "Experience with Snowflake or Redshift" -> Snowflake, Redshift
                words = re.findall(r"\b[A-Z][a-zA-Z0-9\.\+#\-]+\b", line)
                for w in words:
                    w_clean = w.strip(".,;:()\"'")
                    w_lower = w_clean.lower()
                    if w_clean and w_lower not in stop_words_extra and len(w_clean) > 1:
                        # Exclude general capitalized words like "Experience", "Strong", "We"
                        if w_clean not in ["Experience", "Strong", "We", "You", "The", "Must", "Good", "Familiarity", "Understanding", "Excellent", "Working"]:
                            found_skills.add(w_clean)

    extracted["extracted_skills"] = sorted(list(found_skills))

    # 4. Experience match
    exp_match = re.search(r"(\d+[\s–-]+\d+\s*(?:years|yrs))|(\d+\+\s*(?:years|yrs))", text, re.IGNORECASE)
    if exp_match:
        extracted["extracted_experience"] = exp_match.group(0).strip()
    else:
        extracted["extracted_experience"] = "5-9 years"
        
    # 5. Education match
    edu_keywords = ["computer science", "engineering", "b.tech", "m.tech", "ms", "phd", "bachelor", "master"]
    edu_found = []
    for kw in edu_keywords:
        if kw in text.lower():
            edu_found.append(kw)
    extracted["extracted_education"] = ", ".join(edu_found) if edu_found else "Degree in CS or related field"
    
    # 6. Certifications
    cert_keywords = ["aws", "gcp", "azure", "pmp", "tensorflow developer"]
    for kw in cert_keywords:
        if kw in text.lower():
            extracted["extracted_certifications"].append(kw.upper())
            
    # 7. Responsibilities
    resp_section = False
    for line in lines:
        line_lower = line.lower()
        if any(w in line_lower for w in ["responsibility", "responsibilities", "what you'd actually be doing", "what you will do"]):
            resp_section = True
            continue
        if resp_section:
            if any(w in line_lower for w in ["requirement", "qualification", "what we mean", "skills", "experience"]):
                resp_section = False
            elif line.startswith("-") or line.startswith("*") or line.startswith("•") or (len(line) > 10 and line[0].isupper()):
                extracted["extracted_responsibilities"].append(line.lstrip("-*• ").strip())
                if len(extracted["extracted_responsibilities"]) >= 8:
                    resp_section = False
                    
    if not extracted["extracted_responsibilities"]:
        extracted["extracted_responsibilities"] = [
            "Own the core systems and matching layer of the product.",
            "Audit current implementation and identify architectural improvements.",
            "Design and ship scalable v2 features matching modern requirements.",
            "Establish testing and evaluation frameworks to verify system improvements.",
            "Collaborate with multi-functional teams on end-user workflow solutions."
        ]
        
    return extracted
